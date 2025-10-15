from typing import List
import os
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.config import settings

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
        )
    
    def extract_text(self, file_path: str, file_type: str) -> tuple[str, int]:
        """Extract text from document and return text + page count"""
        if file_type == "pdf":
            return self._extract_from_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return self._extract_from_docx(file_path)
        elif file_type == "txt":
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: str) -> tuple[str, int]:
        reader = PdfReader(file_path)
        page_count = len(reader.pages)
        
        if page_count > settings.MAX_PAGES_PER_DOCUMENT:
            raise ValueError(f"Document exceeds maximum page limit of {settings.MAX_PAGES_PER_DOCUMENT}")
        
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text, page_count
    
    def _extract_from_docx(self, file_path: str) -> tuple[str, int]:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        page_count = len(doc.paragraphs) // 10  # Rough estimate
        return text, max(1, page_count)
    
    def _extract_from_txt(self, file_path: str) -> tuple[str, int]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        page_count = len(text) // 3000  # Rough estimate (3000 chars per page)
        return text, max(1, page_count)
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks"""
        chunks = self.text_splitter.split_text(text)
        return chunks